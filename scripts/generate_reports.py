from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

BASE_DIR = Path(__file__).resolve().parents[1]
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

from src.personality_predictor.config import GROUP_ACCENTS, MODEL_CONFIGS, OUTPUT_DOC_DIR

METRICS_PATH = BASE_DIR / "models" / "metrics.json"


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Georgia"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Georgia"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Georgia"
    styles["Heading 2"].font.size = Pt(12.5)


def add_title(document: Document, title: str, subtitle: str) -> None:
    document.add_paragraph(title, style="Title")
    p = document.add_paragraph()
    run = p.add_run(subtitle)
    run.font.size = Pt(11)
    run.italic = True
    document.add_paragraph("")


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph("")


def add_code_block(document: Document, code_text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    document.add_paragraph("")


def load_metrics() -> dict[str, object]:
    if not METRICS_PATH.exists():
        return {}
    return json.loads(METRICS_PATH.read_text(encoding="utf-8"))


def build_pdr(metrics: dict[str, object]) -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc, "Project Definition Report", "Personality Type Predictor - MBTI Edition")

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "This project delivers a Streamlit web application that predicts an MBTI-style personality type after the user answers 10 quiz questions. "
        "The app combines a quiz-to-dimension scoring layer with a machine learning classification layer trained on the Kaggle MBTI 500 dataset."
    )

    doc.add_heading("2. Objectives", level=1)
    add_bullets(
        doc,
        [
            "Train Decision Tree and KNN classifiers on a real Kaggle MBTI dataset with 16 target labels.",
            "Translate 10 quiz responses into dimension scores and a model-ready persona summary.",
            "Present a colorful Streamlit interface with a landing page, question flow, and results dashboard.",
            "Visualize prediction results using result cards, confidence summaries, and a radar chart.",
            "Document the system so the project can be reproduced, extended, and demonstrated in an academic setting.",
        ],
    )

    doc.add_heading("3. Scope", level=1)
    doc.add_heading("3.1 In Scope", level=2)
    add_bullets(
        doc,
        [
            "Loading, cleaning, and vectorizing MBTI 500 text posts.",
            "Training and saving Decision Tree and KNN pipelines.",
            "Ten-question quiz flow with weighted MBTI dimension mapping.",
            "Prediction results with top type candidates and visual feedback.",
            "Project reports, setup notes, and deliverable packaging.",
        ],
    )
    doc.add_heading("3.2 Out of Scope", level=2)
    add_bullets(
        doc,
        [
            "Clinical personality assessment or psychological diagnosis.",
            "User accounts, cloud deployment, and persistent user histories.",
            "Deep learning, transformer fine-tuning, or social media scraping.",
            "Production-scale security hardening or enterprise analytics.",
        ],
    )

    rows_used = str(metrics.get("rows_used", "106067"))
    doc.add_heading("4. Dataset Description", level=1)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Dataset name", "MBTI 500 (Kaggle MBTI text dataset)"],
            ["Source format", "CSV"],
            ["Main columns", "posts, type"],
            ["Observed class count", "16 MBTI labels"],
            ["Rows used", rows_used],
            ["Why it fits", "It links free-text behavior cues to MBTI labels, which supports TF-IDF based text classification."],
        ],
    )

    doc.add_heading("5. Functional Requirements", level=1)
    add_bullets(
        doc,
        [
            "FR-01: Load the MBTI dataset from the workspace data folder or configured fallback path.",
            "FR-02: Clean raw posts by removing URLs, punctuation noise, and direct MBTI label leakage.",
            "FR-03: Train both Decision Tree and KNN models using TF-IDF features and a stratified train/test split.",
            "FR-04: Save trained pipelines as reusable joblib artifacts for app inference.",
            "FR-05: Present 10 quiz questions one at a time with a visible progress bar.",
            "FR-06: Convert quiz answers into MBTI dimension scores and a natural-language persona summary.",
            "FR-07: Let the user pick Decision Tree or KNN before starting the quiz.",
            "FR-08: Show predicted type, confidence, top candidates, and a radar chart on the results screen.",
        ],
    )

    doc.add_heading("6. Non-Functional Requirements", level=1)
    add_bullets(
        doc,
        [
            "NFR-01: The interface should remain readable and colorful on desktop and mobile widths.",
            "NFR-02: The app should respond to quiz clicks without perceptible lag on a local machine.",
            "NFR-03: The codebase should stay modular so training, quiz logic, and reporting can evolve independently.",
            "NFR-04: Model training should be reproducible through command-line scripts and fixed hyperparameters.",
            "NFR-05: Educational disclaimers must state that the project is descriptive and not diagnostic.",
        ],
    )

    doc.add_heading("7. Deliverables", level=1)
    add_bullets(
        doc,
        [
            "Streamlit application source code (app.py, package modules, CSS).",
            "Training script and generated model artifacts.",
            "Project Definition Report, Technical Document, and Design Document in DOCX format.",
            "Requirements file and README with setup/run steps.",
            "Evaluation metrics JSON generated after training.",
        ],
    )

    doc.add_heading("8. Five-Phase Timeline", level=1)
    add_table(
        doc,
        ["Phase", "Focus", "Planned Output"],
        [
            ["Phase 1", "Problem framing and dataset study", "Project scope, dataset understanding, success criteria"],
            ["Phase 2", "Preprocessing and feature engineering", "Cleaned text corpus and TF-IDF configuration"],
            ["Phase 3", "Model training and comparison", "Decision Tree and KNN metrics, saved pipelines"],
            ["Phase 4", "Streamlit UI implementation", "Landing, quiz, and results views with charts"],
            ["Phase 5", "Testing and documentation", "Reports, run guide, and final demonstration package"],
        ],
    )

    doc.add_heading("9. Risk Table", level=1)
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["Class imbalance across MBTI labels", "Some rare types may be predicted less accurately", "Use stratified splitting, balanced tree weights, and transparent confidence display"],
            ["KNN runtime on large sparse vectors", "Training and inference can become slower", "Use sampled training size for KNN and a bounded TF-IDF vocabulary"],
            ["Quiz answers may not perfectly mirror free-text posts", "Model confidence may fluctuate", "Generate richer persona summaries and show both quiz signal and model output"],
            ["Dependency mismatch on a new machine", "The app may fail to launch", "Provide pinned requirements and explicit setup commands"],
            ["Over-interpretation of personality results", "Users may treat the result as a diagnosis", "Place educational disclaimers in the UI and documentation"],
        ],
    )
    return doc


def build_technical_doc(metrics: dict[str, object]) -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc, "Technical Document", "Personality Type Predictor - Implementation Details")

    doc.add_heading("1. Two-Layer Architecture", level=1)
    add_table(
        doc,
        ["Layer", "Responsibilities", "Main Files"],
        [
            ["Presentation layer", "Landing page, quiz navigation, result rendering, custom CSS, chart display", "app.py, assets/styles.css"],
            ["Intelligence layer", "Dataset loading, preprocessing, TF-IDF, model training, quiz scoring, prediction helpers", "train_model.py, src/personality_predictor/*.py"],
        ],
    )

    doc.add_heading("2. File Structure", level=1)
    doc.add_paragraph(
        "New project/\n"
        "  app.py\n"
        "  train_model.py\n"
        "  requirements.txt\n"
        "  README.md\n"
        "  assets/styles.css\n"
        "  models/\n"
        "  output/doc/\n"
        "  scripts/generate_reports.py\n"
        "  src/personality_predictor/config.py\n"
        "  src/personality_predictor/quiz.py\n"
        "  src/personality_predictor/ml.py\n"
        "  src/personality_predictor/charts.py"
    )

    doc.add_heading("3. Data Pipeline", level=1)
    add_bullets(
        doc,
        [
            "Step 1 - Preprocessing: load posts and type, remove missing rows, lowercase text, strip URLs, replace MBTI tokens, and remove punctuation noise.",
            "Step 2 - TF-IDF: convert cleaned posts into sparse unigram/bigram features with bounded vocabulary size.",
            "Step 3 - Split: apply a stratified 80/20 train/test split so every MBTI label appears in both sets.",
            "Step 4 - Model fit: train Decision Tree and KNN pipelines, then evaluate on the held-out test set.",
            "Step 5 - Persistence: save trained pipelines and metrics JSON in the models folder.",
        ],
    )

    doc.add_heading("4. Model Hyperparameters", level=1)
    decision_tree = MODEL_CONFIGS["Decision Tree"]
    knn = MODEL_CONFIGS["KNN"]
    add_table(
        doc,
        ["Decision Tree Parameter", "Value"],
        [
            ["sample_size", str(decision_tree["sample_size"])],
            ["max_features", str(decision_tree["max_features"])],
            ["ngram_range", str(decision_tree["ngram_range"])],
            ["min_df", str(decision_tree["min_df"])],
            ["max_depth", str(decision_tree["max_depth"])],
            ["min_samples_split", str(decision_tree["min_samples_split"])],
            ["min_samples_leaf", str(decision_tree["min_samples_leaf"])],
            ["class_weight", str(decision_tree["class_weight"])],
        ],
    )
    add_table(
        doc,
        ["KNN Parameter", "Value"],
        [
            ["sample_size", str(knn["sample_size"])],
            ["max_features", str(knn["max_features"])],
            ["ngram_range", str(knn["ngram_range"])],
            ["min_df", str(knn["min_df"])],
            ["n_neighbors", str(knn["n_neighbors"])],
            ["weights", str(knn["weights"])],
            ["metric", str(knn["metric"])],
            ["algorithm", str(knn["algorithm"])],
        ],
    )

    doc.add_heading("5. Quiz-to-Dimension Mapping Logic", level=1)
    doc.add_paragraph(
        "Each quiz question maps to one MBTI dimension: IE, SN, TF, or JP. Answers award weighted points to the matching letter. "
        "The dominant letter per dimension forms the quiz signal, such as INTJ. Those answers also become a persona summary sentence block, "
        "which serves as the input text for the selected machine learning pipeline."
    )

    doc.add_heading("6. Key Code Snippets", level=1)
    doc.add_heading("6.1 train_model.py", level=2)
    add_code_block(
        doc,
        "def main():\n"
        "    args = parse_args()\n"
        "    metrics = train_and_save_models(dataset_path=args.data, max_rows=args.max_rows)\n"
        "    print(json.dumps(metrics, indent=2))"
    )
    doc.add_heading("6.2 app.py", level=2)
    add_code_block(
        doc,
        "def finalize_results():\n"
        "    scored = score_answers(st.session_state.answers)\n"
        "    persona_text = compose_persona_text(st.session_state.answers, scored)\n"
        "    prediction = predict_profile(st.session_state.selected_model, persona_text)\n"
        "    st.session_state.result = scored"
    )

    if metrics.get("models"):
        doc.add_heading("7. Current Evaluation Snapshot", level=1)
        rows = []
        for model_name, values in metrics["models"].items():
            rows.append(
                [
                    model_name,
                    str(values["sample_size"]),
                    str(values["accuracy"]),
                    str(values["macro_avg_f1"]),
                    str(values["weighted_avg_f1"]),
                ]
            )
        add_table(doc, ["Model", "Train Sample", "Accuracy", "Macro F1", "Weighted F1"], rows)

    doc.add_heading("8. Dependencies", level=1)
    add_bullets(
        doc,
        [
            "streamlit",
            "pandas",
            "numpy",
            "scikit-learn",
            "joblib",
            "matplotlib",
            "python-docx",
        ],
    )

    doc.add_heading("9. Setup and Run Instructions", level=1)
    add_bullets(
        doc,
        [
            "Install dependencies with python -m pip install -r requirements.txt.",
            "Place MBTI 500.csv inside data/ or use the fallback path already configured in the project.",
            "Train the models with python train_model.py.",
            "Start the app with streamlit run app.py.",
            "Open the local Streamlit URL in a browser and choose Decision Tree or KNN before starting the quiz.",
        ],
    )
    return doc


def build_design_doc() -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc, "Design Document", "Personality Type Predictor - Visual and Interaction Design")

    doc.add_heading("1. Color System", level=1)
    add_table(
        doc,
        ["Token", "Hex", "Usage"],
        [
            ["Paper", "#F7F1E8", "Primary page background"],
            ["Paper Alt", "#FFF9F1", "Cards, quiz surfaces, result tiles"],
            ["Ink", "#152238", "Headlines and primary text"],
            ["Muted", "#5F6B7A", "Secondary copy and captions"],
            ["Analysts Accent", GROUP_ACCENTS["Analysts"], "INTJ, INTP, ENTJ, ENTP results"],
            ["Diplomats Accent", GROUP_ACCENTS["Diplomats"], "INFJ, INFP, ENFJ, ENFP results"],
            ["Sentinels Accent", GROUP_ACCENTS["Sentinels"], "ISTJ, ISFJ, ESTJ, ESFJ results"],
            ["Explorers Accent", GROUP_ACCENTS["Explorers"], "ISTP, ISFP, ESTP, ESFP results"],
        ],
    )

    doc.add_heading("2. Typography Scale", level=1)
    add_table(
        doc,
        ["Level", "Font", "Purpose"],
        [
            ["Display", "Georgia 52 px", "Landing hero and result type headline"],
            ["Heading 1", "Georgia 35 px", "Section titles and screen anchors"],
            ["Heading 2", "Georgia 24 px", "Question prompts and card headings"],
            ["Body", "Trebuchet MS 16 px", "General copy"],
            ["Caption", "Trebuchet MS 14 px", "Labels, helper text, and metadata"],
        ],
    )

    doc.add_heading("3. Screen Flows", level=1)
    add_bullets(
        doc,
        [
            "Landing -> The user reads the project summary, sees the four MBTI groups, chooses a model, and starts the quiz.",
            "Quiz -> One question is shown at a time with a progress bar, a prompt card, and two large answer buttons.",
            "Results -> The predicted type, group card, radar chart, dimension bars, and top candidate types are displayed together.",
        ],
    )

    doc.add_heading("4. Component Specifications", level=1)
    add_table(
        doc,
        ["Component", "Spec"],
        [
            ["Progress bar", "12 px rounded track with warm-to-cool gradient fill and percentage label"],
            ["Question card", "Large serif prompt, helper copy, pill chip for MBTI dimension, soft paper card background"],
            ["Answer buttons", "Two-column layout, full-width gradient buttons, paired with descriptive answer labels"],
            ["Radar chart", "Four-axis polar chart rendered with matplotlib and tinted by MBTI group accent color"],
            ["Result card", "Accent-tinted hero block plus summary card showing model, confidence, quiz signal, and persona text"],
        ],
    )

    doc.add_heading("5. Custom CSS Strategy", level=1)
    doc.add_paragraph(
        "The interface uses a single custom stylesheet injected into Streamlit. CSS variables define reusable design tokens, "
        "while semantic classes style the hero section, group cards, progress bars, question card, dimension rows, and result tiles. "
        "Global button overrides create a more branded look than the Streamlit defaults."
    )

    doc.add_heading("6. Accessibility Notes", level=1)
    add_bullets(
        doc,
        [
            "High contrast is preserved between dark text and light paper backgrounds.",
            "Buttons use generous size and spacing for easier tapping and keyboard focus.",
            "Progress state is communicated with both width change and text labels.",
            "Sections are kept short and scannable to reduce cognitive overload.",
            "The app includes a non-diagnostic disclaimer to reduce misleading interpretation.",
        ],
    )

    doc.add_heading("7. Responsive Layout Approach", level=1)
    doc.add_paragraph(
        "Desktop layouts use two-column hero and results compositions. A media query collapses those grids into a single column below 900 px so the quiz and results remain readable on tablets and phones. "
        "Typography scales down gracefully while buttons continue using the full available width."
    )
    return doc


def main() -> None:
    OUTPUT_DOC_DIR.mkdir(parents=True, exist_ok=True)
    metrics = load_metrics()

    pdr_path = OUTPUT_DOC_DIR / "PDR_Personality_Predictor.docx"
    technical_path = OUTPUT_DOC_DIR / "Technical_Doc_Personality_Predictor.docx"
    design_path = OUTPUT_DOC_DIR / "Design_Doc_Personality_Predictor.docx"

    build_pdr(metrics).save(pdr_path)
    build_technical_doc(metrics).save(technical_path)
    build_design_doc().save(design_path)

    print("Generated:")
    print(pdr_path)
    print(technical_path)
    print(design_path)


if __name__ == "__main__":
    main()
